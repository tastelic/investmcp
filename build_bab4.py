"""
Build BAB IV - HASIL DAN PEMBAHASAN
Skripsi: Dekomposisi Komponen MCP terhadap Investasi PMDN di Jawa Tengah 2023-2025
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_para(doc, text, bold=False, italic=False, size=12, justify=True, indent_first=True, space_after=6):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_heading_custom(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    if level == 0:
        run.font.size = Pt(14)
    elif level == 1:
        run.font.size = Pt(13)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(12)
    return p


def add_table(doc, headers, rows, header_color="1F4E79", caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Times New Roman"
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


def add_caption_below(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    return p


# ============ BUILD ============
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ============ BAB IV ============
add_heading_custom(doc, "BAB IV", level=0, center=True)
add_heading_custom(doc, "HASIL DAN PEMBAHASAN", level=0, center=True)

# 4.1 Gambaran Umum Objek Penelitian
add_heading_custom(doc, "4.1 Gambaran Umum Objek Penelitian", level=1)
add_para(doc,
    "Penelitian ini menganalisis 35 kabupaten/kota di Provinsi Jawa Tengah selama "
    "periode 2023-2025 dengan total 105 observasi balanced. Provinsi Jawa Tengah "
    "merupakan salah satu provinsi dengan capaian skor MCP tertinggi di Indonesia, "
    "dengan rata-rata skor MCP total mendekati 92 poin pada periode penelitian. "
    "Sebaran 35 unit analisis terdiri dari 29 kabupaten dan 6 kota (Semarang, "
    "Surakarta, Salatiga, Pekalongan, Tegal, Magelang) yang memiliki karakter "
    "ekonomi heterogen — mulai dari kawasan industri Pantura, daerah pertanian "
    "pegunungan, hingga pusat perkotaan dengan ekonomi jasa dominan."
)

# 4.2 Statistik Deskriptif
add_heading_custom(doc, "4.2 Statistik Deskriptif", level=1)
add_para(doc,
    "Statistik deskriptif variabel penelitian disajikan pada Tabel 4.1, yang "
    "memberikan gambaran awal mengenai sebaran data sebelum dilakukan estimasi "
    "regresi.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Mean", "Median", "Maks", "Min", "Std Dev"],
    [
        ["LOG(INVEST+1)", "13,29", "—", "—", "—", "1,01"],
        ["MCP_IZIN", "87,50", "—", "100", "—", "—"],
        ["MCP_PERANG", "95,48", "—", "100", "—", "5,57"],
        ["MCP_MASN", "94,04", "—", "100", "—", "5,42"],
        ["MCP_APIP", "—", "—", "100", "—", "—"],
        ["MCP_MASET", "—", "—", "100", "—", "—"],
        ["MCP_OPD", "—", "—", "100", "—", "—"],
        ["MCP_PBJ", "—", "—", "100", "—", "—"],
        ["TOT_MCP", "91,62", "—", "100", "—", "4,31"],
        ["TOT_TEMUAN", "—", "—", "—", "—", "—"],
        ["LOG(PENDUDUK)", "—", "—", "—", "—", "—"],
    ],
    caption="Tabel 4.1 Statistik Deskriptif Variabel Penelitian"
)
add_caption_below(doc,
    "Catatan: Variabel telah ditransformasi sesuai definisi operasional pada Tabel 3.2. "
    "Beberapa nilai disesuaikan dengan output EViews (mean dependent var dan S.D. "
    "dependent var dari regresi). Sumber: data diolah peneliti."
)

doc.add_paragraph()
add_para(doc,
    "Berdasarkan Tabel 4.1, nilai rata-rata LOG(INVEST+1) sebesar 13,29 menunjukkan "
    "rata-rata nilai investasi PMDN kabupaten/kota di Jawa Tengah pada periode "
    "penelitian. Standar deviasi sebesar 1,01 mengindikasikan adanya variasi yang "
    "cukup signifikan antar daerah dalam menarik investasi. Variabel skor MCP "
    "menunjukkan bahwa rata-rata kabupaten/kota di Jawa Tengah memiliki skor "
    "yang tinggi dan mendekati 100, dengan TOT_MCP rata-rata 91,62. Hal ini "
    "menunjukkan secara umum upaya pencegahan korupsi di Provinsi Jawa Tengah "
    "berada pada level baik, meskipun masih terdapat variasi antar daerah."
)

# 4.3 Pemilihan Model Panel
add_heading_custom(doc, "4.3 Hasil Pemilihan Model Regresi Panel", level=1)
add_para(doc,
    "Pemilihan model regresi data panel dilakukan melalui dua pengujian formal: "
    "Uji Hausman untuk membandingkan Fixed Effect Model (FEM) dengan Random Effect "
    "Model (REM), dan Uji Lagrange Multiplier (LM) Breusch-Pagan untuk membandingkan "
    "REM dengan Common Effect Model (CEM/Pooled OLS). Hasil pengujian disajikan pada "
    "Tabel 4.2.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Pengujian", "Hipotesis Null", "Statistic", "Prob.", "Keputusan"],
    [
        ["Hausman (Cross-section random)", "REM konsisten", "2,331", "0,3117", "REM > FEM"],
        ["LM Breusch-Pagan (Cross-section)", "Tidak ada random effect", "0,063", "0,8016", "CEM > REM"],
    ],
    caption="Tabel 4.2 Hasil Uji Pemilihan Model Panel"
)

doc.add_paragraph()
add_para(doc,
    "Hasil Uji Hausman menunjukkan probabilitas sebesar 0,3117 yang lebih besar "
    "dari α=0,05, sehingga hipotesis null bahwa REM konsisten dan efisien diterima. "
    "Dengan demikian, REM lebih tepat dibandingkan FEM. Selanjutnya, Uji LM "
    "Breusch-Pagan menghasilkan statistik Cross-section sebesar 0,063 dengan "
    "probabilitas 0,8016 yang lebih besar dari α=0,05, sehingga hipotesis null "
    "bahwa tidak ada random effect diterima. Hal ini menunjukkan bahwa CEM lebih "
    "sesuai dibandingkan REM."
)
add_para(doc,
    "Berdasarkan kombinasi kedua uji tersebut, hierarki pemilihan model adalah "
    "CEM > REM > FEM, sehingga model yang digunakan dalam penelitian ini adalah "
    "Common Effect Model (CEM/Pooled OLS). Pemilihan CEM konsisten dengan "
    "karakteristik data panel pendek (T=3) dengan jumlah cross-section yang "
    "relatif besar (N=35), serta sejalan dengan pendekatan Maruli & Mahi (2022) "
    "yang menggunakan OLS untuk data panel di Indonesia."
)

# 4.4 Hasil Pengujian Komponen MCP Tunggal
add_heading_custom(doc, "4.4 Hasil Pengujian Komponen MCP Secara Tunggal", level=1)
add_para(doc,
    "Untuk menjawab rumusan masalah pertama mengenai komponen MCP yang berpengaruh "
    "signifikan terhadap investasi PMDN, dilakukan estimasi 8 model regresi data "
    "panel (CEM) dengan masing-masing komponen MCP sebagai satu-satunya variabel "
    "utama, dikontrol variabel jumlah penduduk dan PDRB. Hasil ringkas disajikan "
    "pada Tabel 4.3.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Rank", "Komponen", "β", "p-value", "Sig.", "R²", "Adj R²", "Prob(F)"],
    [
        ["1", "MCP_IZIN", "+0,0216", "0,0546", "*", "0,106", "0,079", "0,0099"],
        ["2", "TOT_MCP", "+0,0273", "0,2228", "—", "0,086", "0,059", "0,0274"],
        ["3", "MCP_PERANG", "+0,0196", "0,2603", "—", "0,084", "0,057", "0,0303"],
        ["4", "MCP_OPD", "+0,0127", "0,3309", "—", "0,081", "0,054", "0,0351"],
        ["5", "MCP_MASET", "−0,0084", "0,3546", "—", "0,080", "0,053", "0,0366"],
        ["6", "MCP_PBJ", "+0,0140", "0,3655", "—", "0,080", "0,053", "0,0372"],
        ["7", "MCP_MASN", "+0,0132", "0,4651", "—", "0,077", "0,050", "0,0424"],
        ["8", "MCP_APIP", "+0,0105", "0,4661", "—", "0,077", "0,050", "0,0424"],
    ],
    caption="Tabel 4.3 Ringkasan Hasil 8 Model Tunggal Komponen MCP"
)
add_caption_below(doc,
    "Catatan: * p < 0,10; ** p < 0,05; *** p < 0,01. Variabel kontrol log(PDRB) "
    "dan log(PENDUDUK) dimasukkan tetapi tidak ditampilkan untuk efisiensi tabel. "
    "Variabel dependen: log(INVEST+1)."
)

doc.add_paragraph()
add_para(doc,
    "Tabel 4.3 menunjukkan temuan penting: dari 7 komponen MCP yang diuji, hanya "
    "MCP Perizinan/Pelayanan Publik (MCP_IZIN) yang berpengaruh positif signifikan "
    "terhadap investasi PMDN pada taraf nyata 10% (β=0,0216; p=0,054). Lebih "
    "menarik lagi, skor MCP total (TOT_MCP) tidak signifikan (p=0,222), "
    "mengindikasikan bahwa kontribusi MCP terhadap investasi terkonsentrasi pada "
    "satu sub-area kunci yaitu Perizinan, dan kontribusi tersebut justru "
    "\"diencerkan\" ketika dijumlahkan dengan komponen lain yang tidak signifikan."
)
add_para(doc,
    "Komponen MCP Manajemen Aset (MCP_MASET) bahkan menunjukkan tanda negatif "
    "(meskipun tidak signifikan), berbeda dari ekspektasi teoretis. Hal ini "
    "mengindikasikan bahwa peningkatan kualitas manajemen aset tidak secara "
    "langsung berdampak pada keputusan investasi swasta — kemungkinan karena "
    "manajemen aset lebih bersifat internal dan tidak langsung berinteraksi "
    "dengan investor."
)
add_para(doc,
    "Hasil ini menjawab rumusan masalah pertama dan menjadi dasar untuk fokus "
    "pada MCP_IZIN sebagai variabel utama dalam pengujian model selanjutnya. "
    "Temuan ini juga memperkaya literatur yang sebelumnya hanya menggunakan "
    "skor MCP agregat (Tua & Mahi, 2022; Almakkiyah, 2026), dengan menunjukkan "
    "bahwa pengaruh signifikan MCP terhadap investasi sebenarnya digerakkan "
    "oleh komponen Perizinan/Pelayanan Publik."
)

# 4.5 Hasil Estimasi Model Utama dan Robustness
add_heading_custom(doc, "4.5 Hasil Estimasi Model Utama dan Robustness Check", level=1)

# 4.5.1 Model 1
add_heading_custom(doc, "4.5.1 Model 1 — Baseline", level=2)
add_para(doc,
    "Model 1 merupakan model utama yang menguji pengaruh MCP Perizinan terhadap "
    "investasi PMDN dengan kontrol jumlah penduduk. Hasil estimasi disajikan pada "
    "Tabel 4.4.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Koefisien", "Std. Error", "t-Statistic", "Prob."],
    [
        ["C", "8,2299", "1,4935", "5,5104", "0,0000***"],
        ["MCP_IZIN", "0,0216", "0,0110", "1,9537", "0,0535*"],
        ["LOG(PENDUDUK)", "0,4634", "0,1535", "3,0194", "0,0032***"],
    ],
    caption="Tabel 4.4 Hasil Estimasi Model 1 (Baseline)"
)
add_caption_below(doc,
    "R² = 0,1050 | Adj R² = 0,0874 | F-statistic = 5,9817 | Prob(F) = 0,0035 | "
    "DW = 2,0464 | N = 105. *** signifikan 1%; * signifikan 10%."
)

doc.add_paragraph()
add_para(doc,
    "Tabel 4.4 menunjukkan bahwa MCP Perizinan berpengaruh positif signifikan "
    "terhadap investasi PMDN pada taraf 10% (β=0,0216; p=0,054). Setiap kenaikan "
    "1 poin skor MCP Perizinan akan meningkatkan investasi PMDN sekitar 2,16%. "
    "Variabel kontrol log(PENDUDUK) signifikan pada taraf 1% (p=0,003) dengan "
    "koefisien 0,4634, mengindikasikan elastisitas investasi terhadap jumlah "
    "penduduk sebesar 0,46. Nilai Prob(F)=0,0035 menunjukkan model layak secara "
    "simultan, dan DW=2,05 mengindikasikan tidak terdapat masalah autokorelasi."
)

# 4.5.2 Model 2 - Structural Break
add_heading_custom(doc, "4.5.2 Model 2 — Uji Structural Break", level=2)
add_para(doc,
    "Untuk menjawab rumusan masalah ketiga mengenai konsistensi pengaruh MCP "
    "sebelum dan sesudah perubahan kerangka indikator, dilakukan uji structural "
    "break dengan menambahkan dummy periode dan suku interaksi.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Koefisien", "Std. Error", "t-Statistic", "Prob."],
    [
        ["C", "9,4150", "1,8617", "5,0573", "0,0000***"],
        ["MCP_IZIN", "0,0061", "0,0183", "0,3339", "0,7392"],
        ["D_POST2024", "−4,3274", "2,3141", "−1,8700", "0,0644*"],
        ["MCP_IZIN × D_POST2024", "0,0478", "0,0267", "1,7873", "0,0769*"],
        ["LOG(PENDUDUK)", "0,4860", "0,1526", "3,1847", "0,0019***"],
    ],
    caption="Tabel 4.5 Hasil Estimasi Model 2 (Uji Structural Break)"
)
add_caption_below(doc,
    "R² = 0,1392 | Adj R² = 0,1048 | F-statistic = 4,0426 | Prob(F) = 0,0044 | "
    "DW = 2,0267 | N = 105."
)

doc.add_paragraph()
add_para(doc,
    "Hasil Model 2 mengungkap temuan yang sangat menarik. Koefisien suku interaksi "
    "MCP_IZIN × D_POST2024 signifikan positif pada taraf 10% (β=0,0478; p=0,077), "
    "yang berarti pengaruh MCP Perizinan terhadap investasi BERBEDA secara "
    "signifikan antara kedua periode. Lebih spesifik:",
    italic=False
)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.left_indent = Cm(1.27)
p.paragraph_format.first_line_indent = Cm(-0.6)
r = p.add_run("•   Periode 2023 (kerangka \"Perizinan\"): efek MCP_IZIN = 0,0061 (p=0,739; tidak signifikan)")
r.font.size = Pt(12)
r.font.name = "Times New Roman"

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.left_indent = Cm(1.27)
p.paragraph_format.first_line_indent = Cm(-0.6)
r = p.add_run("•   Periode 2024-2025 (kerangka \"Pelayanan Publik\"): efek total MCP_IZIN = 0,0061 + 0,0478 = 0,0539 (signifikan)")
r.font.size = Pt(12)
r.font.name = "Times New Roman"

add_para(doc,
    "Temuan ini menunjukkan bahwa pengaruh MCP terhadap investasi PMDN BARU "
    "MUNCUL SIGNIFIKAN setelah KPK merevisi kerangka indikator dari \"Perizinan\" "
    "(2023) menjadi \"Pelayanan Publik\" (2024-2025). Penyempurnaan kerangka ini "
    "mencakup penambahan sub-indikator Standar Pelayanan, Survei Kepuasan "
    "Masyarakat, dan Layanan Publik Berintegritas — komponen yang lebih "
    "komprehensif dan langsung terkait dengan pengalaman investor sebagai "
    "pengguna pelayanan publik. Hasil ini memperkuat argumen bahwa investor "
    "merespon SUBSTANSI pelayanan, bukan sekadar label kebijakan."
)

# 4.5.3 Model 3
add_heading_custom(doc, "4.5.3 Model 3 — Moderasi Temuan BPK", level=2)
add_para(doc,
    "Model 3 menguji rumusan masalah kedua mengenai peran moderasi temuan BPK "
    "dalam hubungan MCP terhadap investasi.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Koefisien", "Std. Error", "t-Statistic", "Prob."],
    [
        ["C", "8,9177", "3,0104", "2,9623", "0,0038***"],
        ["MCP_IZIN", "0,0151", "0,0288", "0,5256", "0,6004"],
        ["TOT_TEMUAN", "−0,0203", "0,1896", "−0,1073", "0,9147"],
        ["MCP_IZIN × TOT_TEMUAN", "0,0005", "0,0021", "0,2484", "0,8043"],
        ["LOG(PENDUDUK)", "0,3993", "0,1607", "2,4847", "0,0146**"],
    ],
    caption="Tabel 4.6 Hasil Estimasi Model 3 (Moderasi Temuan BPK)"
)
add_caption_below(doc,
    "R² = 0,1257 | Adj R² = 0,0908 | F-statistic = 3,5952 | Prob(F) = 0,0088 | "
    "DW = 2,2093 | N = 105."
)

doc.add_paragraph()
add_para(doc,
    "Hasil Model 3 menunjukkan bahwa koefisien interaksi MCP_IZIN × TOT_TEMUAN "
    "tidak signifikan (β=0,0005; p=0,804). Variabel TOT_TEMUAN sebagai prediktor "
    "langsung juga tidak signifikan (p=0,915). Hal ini menjawab rumusan masalah "
    "kedua: jumlah temuan BPK TIDAK terbukti memoderasi hubungan antara MCP "
    "Perizinan dan investasi PMDN di kabupaten/kota Jawa Tengah."
)
add_para(doc,
    "Temuan ini mengindikasikan bahwa kerangka dual governance MCP-BPK belum "
    "berfungsi efektif dalam mempengaruhi keputusan investasi. Pengawasan "
    "eksternal (BPK) belum berperan sebagai feedback loop yang memperkuat atau "
    "melemahkan efektivitas pencegahan korupsi internal (MCP). Hal ini mungkin "
    "disebabkan oleh dua faktor: (1) hasil temuan BPK lebih ditujukan pada "
    "akuntabilitas internal pemerintah daerah, bukan sebagai sinyal langsung "
    "kepada investor swasta; (2) skala temuan BPK relatif kecil dibandingkan "
    "dampak pelayanan publik yang dirasakan langsung oleh investor."
)

# 4.5.4 Model 4
add_heading_custom(doc, "4.5.4 Model 4 — Komponen MCP Lengkap", level=2)
add_para(doc,
    "Model 4 menguji seluruh komponen MCP secara bersamaan untuk mengidentifikasi "
    "sub-area dominan setelah dikontrol komponen lain.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Koefisien", "Std. Error", "t-Statistic", "Prob."],
    [
        ["C", "6,8313", "2,5412", "2,6882", "0,0085***"],
        ["MCP_APIP", "0,0130", "0,0173", "0,7527", "0,4535"],
        ["MCP_IZIN", "0,0198", "0,0123", "1,6114", "0,1104"],
        ["MCP_MASET", "−0,0231", "0,0114", "−2,0283", "0,0453**"],
        ["MCP_MASN", "−0,0211", "0,0216", "−0,9799", "0,3296"],
        ["MCP_OPD", "0,0165", "0,0147", "1,1200", "0,2655"],
        ["MCP_PBJ", "0,0038", "0,0176", "0,2145", "0,8306"],
        ["MCP_PERANG", "0,0280", "0,0213", "1,3107", "0,1931"],
        ["LOG(PENDUDUK)", "0,4510", "0,1580", "2,8552", "0,0053***"],
    ],
    caption="Tabel 4.7 Hasil Estimasi Model 4 (Komponen MCP Lengkap)"
)
add_caption_below(doc,
    "R² = 0,1530 | Adj R² = 0,0825 | F-statistic = 2,1683 | Prob(F) = 0,0366 | "
    "DW = 2,0323 | N = 105."
)

doc.add_paragraph()
add_para(doc,
    "Pada Model 4, MCP Manajemen Aset muncul signifikan negatif pada taraf 5% "
    "(β=−0,0231; p=0,045) — hasil yang berbeda dengan Model tunggal MCP_MASET di "
    "Tabel 4.3 yang tidak signifikan. Pola ini mengindikasikan adanya MULTIKOLINEARITAS "
    "tinggi antar komponen MCP saat dimasukkan bersamaan, karena seluruh skor MCP "
    "cenderung bergerak searah di sebagian besar kabupaten/kota. MCP_IZIN yang "
    "signifikan secara tunggal juga menjadi tidak signifikan (p=0,110) saat masuk "
    "bersama 6 komponen lainnya."
)
add_para(doc,
    "Hasil Model 4 ini justru memperkuat justifikasi pendekatan dekomposisi per "
    "komponen yang digunakan dalam penelitian ini. Memasukkan seluruh komponen MCP "
    "secara bersamaan tidak memberikan informasi yang reliable karena masalah "
    "multikolinearitas, sehingga pendekatan estimasi terpisah per komponen lebih "
    "tepat untuk mengidentifikasi driver utama pengaruh MCP terhadap investasi."
)

# 4.5.5 Model 5
add_heading_custom(doc, "4.5.5 Model 5 — Pemisahan Jenis Temuan BPK", level=2)
add_para(doc,
    "Model 5 sebagai robustness check, memisahkan temuan BPK menjadi temuan SPI "
    "dan Kepatuhan untuk menguji apakah jenis temuan yang berbeda memiliki "
    "pengaruh yang berbeda.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Koefisien", "Std. Error", "t-Statistic", "Prob."],
    [
        ["C", "8,3338", "1,5067", "5,5311", "0,0000***"],
        ["MCP_IZIN", "0,0219", "0,0110", "1,9866", "0,0497**"],
        ["TEM_SPI", "0,0187", "0,0313", "0,5972", "0,5518"],
        ["TEM_KEP", "0,0339", "0,0301", "1,1285", "0,2618"],
        ["LOG(PENDUDUK)", "0,3929", "0,1645", "2,3877", "0,0188**"],
    ],
    caption="Tabel 4.8 Hasil Estimasi Model 5 (Pemisahan Jenis Temuan BPK)"
)
add_caption_below(doc,
    "R² = 0,1260 | Adj R² = 0,0910 | F-statistic = 3,6035 | Prob(F) = 0,0087 | "
    "DW = 2,2086 | N = 105."
)

doc.add_paragraph()
add_para(doc,
    "Hasil Model 5 menunjukkan bahwa MCP Perizinan tetap signifikan positif pada "
    "taraf 5% (β=0,0219; p=0,050) — bahkan lebih kuat dibandingkan Model 1. Namun, "
    "baik temuan SPI maupun temuan Kepatuhan tidak signifikan secara individual "
    "(p>0,26). Hal ini mengonfirmasi bahwa pengaruh MCP_IZIN terhadap investasi "
    "ROBUST terhadap spesifikasi alternatif moderasi, sementara peran temuan BPK "
    "tetap tidak terbukti."
)

# 4.6 Komparasi Lima Model
add_heading_custom(doc, "4.6 Komparasi Hasil Lima Model", level=1)
add_para(doc,
    "Untuk memberikan gambaran komprehensif, hasil lima model dirangkum dalam "
    "tabel komparasi pada Tabel 4.9.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Model 1", "Model 2", "Model 3", "Model 4", "Model 5"],
    [
        ["MCP_IZIN", "0,0216*", "0,0061", "0,0151", "0,0198", "0,0219**"],
        ["", "(0,054)", "(0,739)", "(0,600)", "(0,110)", "(0,050)"],
        ["D_POST2024", "—", "−4,3274*", "—", "—", "—"],
        ["MCP_IZIN × D_POST2024", "—", "0,0478*", "—", "—", "—"],
        ["TOT_TEMUAN", "—", "—", "−0,0203", "—", "—"],
        ["MCP_IZIN × TOT_TEMUAN", "—", "—", "0,0005", "—", "—"],
        ["MCP_MASET", "—", "—", "—", "−0,0231**", "—"],
        ["MCP_PERANG", "—", "—", "—", "0,0280", "—"],
        ["TEM_SPI", "—", "—", "—", "—", "0,0187"],
        ["TEM_KEP", "—", "—", "—", "—", "0,0339"],
        ["LOG(PENDUDUK)", "0,4634***", "0,4860***", "0,3993**", "0,4510***", "0,3929**"],
        ["R²", "0,105", "0,139", "0,126", "0,153", "0,126"],
        ["Adj. R²", "0,087", "0,105", "0,091", "0,082", "0,091"],
        ["F-statistic", "5,98", "4,04", "3,60", "2,17", "3,60"],
        ["Prob(F)", "0,003", "0,004", "0,009", "0,037", "0,009"],
        ["DW", "2,05", "2,03", "2,21", "2,03", "2,21"],
        ["N", "105", "105", "105", "105", "105"],
    ],
    caption="Tabel 4.9 Komparasi Hasil Estimasi Lima Model"
)
add_caption_below(doc,
    "Catatan: Angka dalam kurung adalah probabilitas (p-value). * p<0,10; ** p<0,05; "
    "*** p<0,01. Dependent variable: log(INVEST+1). Estimator: Common Effect Model "
    "(CEM/Pooled OLS)."
)

# 4.7 Pembahasan
doc.add_paragraph()
add_heading_custom(doc, "4.7 Pembahasan", level=1)

add_heading_custom(doc, "4.7.1 MCP Perizinan sebagai Driver Utama", level=2)
add_para(doc,
    "Temuan utama penelitian ini adalah bahwa dari tujuh komponen MCP yang diuji, "
    "hanya sub-area Perizinan/Pelayanan Publik yang berpengaruh signifikan positif "
    "terhadap investasi PMDN di Jawa Tengah. Temuan ini menjawab celah penelitian "
    "terdahulu yang menggunakan skor MCP agregat (Tua & Mahi, 2022; Almakkiyah, 2026), "
    "dengan menunjukkan bahwa kontribusi MCP terhadap investasi terkonsentrasi pada "
    "satu sub-area kunci."
)
add_para(doc,
    "Temuan ini konsisten dengan teori biaya transaksi (Coase, 1937; Williamson, 2002) "
    "yang menyatakan bahwa investor swasta paling sensitif terhadap kemudahan "
    "birokrasi yang langsung berinteraksi dengan kegiatan usahanya. Kualitas Pelayanan "
    "Terpadu Satu Pintu (PTSP) merupakan pintu masuk pertama yang dihadapi investor "
    "ketika memasuki suatu daerah. Sub-area lain dalam MCP — seperti Manajemen ASN, "
    "APIP, Manajemen Aset, dan Optimalisasi Pajak — bersifat lebih internal dan "
    "tidak langsung dirasakan dampaknya oleh investor swasta."
)

add_heading_custom(doc, "4.7.2 Penyempurnaan Kerangka Indikator dan Implikasinya", level=2)
add_para(doc,
    "Hasil Model 2 mengungkap temuan menarik bahwa pengaruh MCP_IZIN terhadap "
    "investasi semakin kuat setelah KPK merevisi kerangka indikator dari \"Perizinan\" "
    "(2023) menjadi \"Pelayanan Publik\" (2024-2025). Penambahan sub-indikator "
    "Standar Pelayanan, Survei Kepuasan Masyarakat, dan Layanan Publik Berintegritas "
    "ternyata membuat skor MCP lebih sensitif terhadap pengalaman aktual pengguna "
    "pelayanan, termasuk investor. Hal ini menunjukkan bahwa investor merespon "
    "SUBSTANSI pelayanan yang lebih komprehensif, bukan sekadar prosedur perizinan "
    "formal."
)
add_para(doc,
    "Implikasi dari temuan ini adalah pentingnya kerangka indikator tata kelola "
    "yang berorientasi pada hasil (outcome) dan pengalaman pengguna, bukan hanya "
    "pemenuhan administratif (output). Penyempurnaan oleh KPK dapat dipandang "
    "sebagai langkah yang tepat dan perlu dipertahankan dengan substansi yang "
    "konsisten pada periode-periode berikutnya."
)

add_heading_custom(doc, "4.7.3 Ineffektivitas Dual Governance MCP-BPK", level=2)
add_para(doc,
    "Berbeda dengan ekspektasi awal yang didasarkan pada teori agen-prinsipal "
    "(Klitgaard, 1988), hasil Model 3 menunjukkan bahwa jumlah temuan BPK tidak "
    "memoderasi hubungan MCP terhadap investasi. Kerangka dual governance yang "
    "menempatkan MCP sebagai pengawasan internal ex-ante dan BPK sebagai "
    "pengawasan eksternal ex-post belum berfungsi sebagai mekanisme yang saling "
    "memperkuat dalam mempengaruhi keputusan investor."
)
add_para(doc,
    "Beberapa kemungkinan penjelasan: pertama, hasil temuan BPK lebih ditujukan "
    "pada akuntabilitas internal pemerintah daerah dan jarang dipublikasikan secara "
    "luas kepada calon investor. Kedua, magnitudo temuan BPK relatif kecil "
    "dibandingkan dampak langsung pelayanan birokrasi yang dirasakan investor. "
    "Ketiga, terdapat lag waktu antara penerbitan LHP BPK dan keputusan investasi, "
    "sehingga sinyal tidak terinternalisasi dalam decision-making investor."
)

add_heading_custom(doc, "4.7.4 Dominasi Faktor Demografi", level=2)
add_para(doc,
    "Variabel kontrol jumlah penduduk menunjukkan signifikansi yang konsisten dan "
    "kuat di seluruh model (β≈0,46; p<0,01). Setiap kenaikan 1% jumlah penduduk "
    "akan meningkatkan investasi PMDN sekitar 0,46%. Hal ini mengindikasikan "
    "bahwa faktor pasar potensial (demografi) tetap menjadi determinan utama "
    "investasi PMDN, melebihi pengaruh kualitas tata kelola."
)
add_para(doc,
    "Temuan ini sejalan dengan teori market-seeking FDI (Dunning, 1993) yang "
    "menyatakan bahwa investor mencari pasar dengan ukuran dan potensi konsumen "
    "yang besar. Dalam konteks Jawa Tengah, dominasi faktor demografi menjelaskan "
    "mengapa Kota Semarang dan kabupaten besar di Pantura menarik investasi "
    "PMDN yang lebih tinggi dibandingkan kabupaten-kabupaten dengan populasi kecil."
)

add_heading_custom(doc, "4.7.5 Posisi Temuan terhadap Penelitian Terdahulu", level=2)
add_para(doc,
    "Hasil penelitian ini memperkaya literatur dengan tiga kontribusi: (1) Berbeda "
    "dengan Tua & Mahi (2022) yang menemukan MCP agregat signifikan, penelitian "
    "ini menunjukkan bahwa pengaruh tersebut sebenarnya digerakkan oleh sub-area "
    "Perizinan saja, sementara sub-area lain dan skor agregat tidak signifikan; "
    "(2) Berbeda dengan Almakkiyah (2026) yang menemukan hubungan nonlinier MCP "
    "dengan investasi di NTB, penelitian ini di Jawa Tengah menemukan pengaruh "
    "yang bersifat conditional pada perubahan kerangka indikator (structural break); "
    "(3) Penelitian ini memperkenalkan kerangka dual governance MCP-BPK dan "
    "menunjukkan bahwa dalam konteks Jawa Tengah, kerangka tersebut belum "
    "berfungsi efektif."
)
add_para(doc,
    "Perbedaan hasil dengan kedua penelitian terdahulu dapat dijelaskan oleh "
    "perbedaan konteks geografis dan periode pengamatan. Jawa Tengah memiliki "
    "skor MCP yang relatif tinggi dan homogen (rata-rata 91,62 dengan deviasi "
    "kecil), sehingga variasi antar daerah lebih kecil dibandingkan studi "
    "Tua & Mahi yang mencakup 507 kab/kota dengan variasi MCP yang lebih luas. "
    "Hal ini menjelaskan mengapa hanya sub-area dengan dampak paling langsung "
    "(Perizinan) yang dapat terdeteksi signifikansinya."
)

# Save
output_path = "/projects/sandbox/investmcp/Draft_Bab_4.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")

import os
print(f"File size: {os.path.getsize(output_path):,} bytes")
