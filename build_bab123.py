"""
Build comprehensive draft Bab I-III untuk skripsi
Pengaruh Komponen MCP terhadap Investasi PMDN di Jawa Tengah 2023-2025
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_para(doc, text, bold=False, italic=False, size=12, justify=True, indent_first=True, space_after=6):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_heading_custom(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    if level == 0:
        run.font.size = Pt(14)
    elif level == 1:
        run.font.size = Pt(13)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(12)
    return p


def add_table(doc, headers, rows, header_color="1F4E79", caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Times New Roman"
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r, row in enumerate(rows):
        cells = table.rows[r + 1].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


def add_equation(doc, text, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    if label:
        p.add_run("    " + label).font.size = Pt(11)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    p.paragraph_format.line_spacing = 1.5
    run = p.runs[0] if p.runs else p.add_run("")
    p.text = ""
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"
    return p


# ============ BUILD ============
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# Page margins (Indonesian skripsi standard)
for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ============ HALAMAN JUDUL ============
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("DEKOMPOSISI KOMPONEN MONITORING CENTER FOR PREVENTION (MCP) "
                          "TERHADAP INVESTASI PENANAMAN MODAL DALAM NEGERI: "
                          "STUDI EMPIRIS DI PROVINSI JAWA TENGAH PERIODE 2023-2025")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Times New Roman"
title_para.paragraph_format.line_spacing = 1.5

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("PROPOSAL/SKRIPSI")
r.bold = True
r.font.size = Pt(13)
r.font.name = "Times New Roman"

doc.add_paragraph()
doc.add_paragraph()

# ============ ABSTRAK ============
add_heading_custom(doc, "ABSTRAK", level=0, center=True)

abstrak = """Investasi merupakan komponen strategis dalam mendorong pertumbuhan ekonomi daerah, sementara kualitas tata kelola pemerintahan menjadi salah satu determinan keputusan investor. Penelitian terdahulu (Munajah, 2026; Maruli & Mahi, 2022) menggunakan skor Monitoring Center for Prevention (MCP) secara agregat sehingga belum dapat mengidentifikasi komponen pencegahan korupsi yang sebenarnya menjadi penggerak hubungan dengan investasi. Penelitian ini melakukan dekomposisi MCP menjadi tujuh sub-area untuk menganalisis pengaruh masing-masing komponen terhadap investasi Penanaman Modal Dalam Negeri (PMDN) di 35 kabupaten/kota Provinsi Jawa Tengah periode 2023-2025 (N=105 observasi balanced). Common Effect Model (CEM/Pooled OLS) dipilih sebagai model regresi data panel berdasarkan Uji Hausman (Prob=0,3117) dan Uji Lagrange Multiplier Breusch-Pagan (Prob=0,8016). Lima spesifikasi model diestimasi termasuk uji structural break terkait perubahan label sub-area MCP. Hasil menunjukkan bahwa dari tujuh komponen MCP yang diuji, hanya MCP Perizinan/Pelayanan Publik yang berpengaruh positif signifikan terhadap investasi PMDN (β=0,0216; p=0,054), dan pengaruh ini semakin kuat (β=0,054) setelah kerangka indikator KPK direvisi menjadi 'Pelayanan Publik' yang lebih komprehensif pada 2024-2025. Variabel temuan BPK tidak terbukti memoderasi hubungan MCP-investasi, sementara faktor demografi (jumlah penduduk) konsisten signifikan di seluruh model. Penelitian ini memberikan kontribusi metodologis berupa dekomposisi MCP per sub-area dan kerangka dual governance MCP-BPK yang belum pernah diuji di literatur sebelumnya. Implikasi kebijakan menunjukkan pemerintah daerah perlu memprioritaskan reformasi Pelayanan Terpadu Satu Pintu (PTSP) dan kualitas pelayanan publik komprehensif sebagai strategi peningkatan investasi."""

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.27)
p.paragraph_format.line_spacing = 1.5
run = p.add_run(abstrak)
run.font.size = Pt(11)
run.font.name = "Times New Roman"

doc.add_paragraph()
kw = doc.add_paragraph()
kw.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r1 = kw.add_run("Kata Kunci: ")
r1.bold = True
r1.italic = True
r1.font.size = Pt(11)
r2 = kw.add_run("Investasi PMDN, Monitoring Center for Prevention, Pencegahan Korupsi, Pelayanan Publik, Temuan BPK, Data Panel Jawa Tengah")
r2.italic = True
r2.font.size = Pt(11)

doc.add_page_break()

# ============ BAB I PENDAHULUAN ============
add_heading_custom(doc, "BAB I", level=0, center=True)
add_heading_custom(doc, "PENDAHULUAN", level=0, center=True)

add_heading_custom(doc, "1.1 Latar Belakang", level=1)

add_para(doc,
    "Investasi merupakan salah satu komponen strategis dalam mendorong pertumbuhan ekonomi daerah karena berperan dalam meningkatkan kapasitas produksi, menciptakan lapangan kerja, serta memperluas aktivitas ekonomi (Almakkiyah, 2026). Daerah dengan arus investasi tinggi cenderung memiliki kemampuan pembangunan yang lebih besar dibandingkan daerah dengan investasi rendah. Oleh karena itu, baik pemerintah pusat maupun daerah terus berupaya menciptakan iklim investasi yang kondusif melalui berbagai kebijakan struktural dan institusional, termasuk peningkatan kualitas tata kelola pemerintahan."
)

add_para(doc,
    "Salah satu faktor yang mempengaruhi keputusan investasi adalah kualitas tata kelola pemerintahan, termasuk upaya pencegahan dan pengendalian korupsi. Teori kelembagaan (North, 1991) menyatakan bahwa investor mempertimbangkan stabilitas kebijakan, transparansi birokrasi, serta jaminan kepastian hukum sebelum memutuskan masuk ke suatu daerah. Praktik korupsi yang tinggi dapat menambah biaya transaksi, meningkatkan ketidakpastian, serta menurunkan kualitas pelayanan publik sehingga menghambat investasi (Mauro, 1995). Sebaliknya, upaya pencegahan korupsi di sektor publik dapat memperkuat kepercayaan investor dan meningkatkan iklim usaha yang lebih sehat."
)

add_para(doc,
    "Di Indonesia, Komisi Pemberantasan Korupsi (KPK) mengembangkan instrumen Monitoring Center for Prevention (MCP) untuk menilai upaya pencegahan korupsi di pemerintah daerah. MCP terdiri dari delapan area tata kelola yang dianggap rawan korupsi: Perencanaan dan Penganggaran APBD, Pengadaan Barang dan Jasa, Pelayanan Terpadu Satu Pintu/Perizinan, Pengawasan APIP, Manajemen ASN, Optimalisasi Pajak Daerah, Manajemen Aset, dan Manajemen Dana Desa. Skor MCP berkisar 0-100, dengan skor mendekati 100 menunjukkan kualitas pencegahan korupsi yang lebih baik (Tua & Mahi, 2022)."
)

add_para(doc,
    "Provinsi Jawa Tengah merupakan salah satu provinsi dengan capaian skor MCP tertinggi di Indonesia, namun realisasi investasi PMDN antar 35 kabupaten/kota di provinsi ini menunjukkan variasi yang signifikan. Beberapa daerah seperti Kota Semarang, Kabupaten Cilacap, dan Kabupaten Kendal memiliki nilai investasi PMDN yang sangat besar, sementara kabupaten lainnya memiliki realisasi yang relatif rendah. Variasi ini menimbulkan pertanyaan mengenai sejauh mana kualitas tata kelola yang diukur melalui MCP berkontribusi terhadap perbedaan kemampuan daerah dalam menarik investasi PMDN."
)

add_para(doc,
    "Penelitian terdahulu mengenai pengaruh MCP terhadap investasi telah dilakukan dalam berbagai konteks. Tua dan Mahi (2022) menggunakan data 507 kabupaten/kota di Indonesia periode 2018-2020 dan menemukan bahwa MCP berpengaruh positif signifikan terhadap investasi swasta. Almakkiyah (2026) menggunakan data 10 kabupaten/kota di Provinsi Nusa Tenggara Barat periode 2019-2024 dan menemukan hubungan nonlinier (threshold consistency) antara MCP dengan investasi. Namun demikian, kedua penelitian tersebut menggunakan skor MCP secara agregat sehingga belum dapat mengidentifikasi sub-area mana dari delapan area MCP yang sebenarnya menjadi penggerak utama hubungan MCP-investasi."
)

add_para(doc,
    "Penggunaan skor MCP agregat memiliki keterbatasan substantif karena berasumsi bahwa seluruh sub-area memiliki kontribusi yang setara terhadap keputusan investor. Padahal, secara teoretis dan empiris, investor swasta cenderung lebih sensitif terhadap aspek pelayanan publik yang langsung berinteraksi dengan kegiatan usahanya, seperti perizinan, dibandingkan dengan aspek tata kelola internal pemerintah daerah seperti manajemen ASN atau pengawasan APIP (Klitgaard, 1988; Camargo, 2011). Oleh karena itu, dekomposisi MCP per sub-area menjadi penting untuk mengidentifikasi prioritas reformasi tata kelola yang lebih tepat sasaran."
)

add_para(doc,
    "Selain MCP sebagai instrumen pengawasan internal, pemerintah daerah juga diawasi oleh Badan Pemeriksa Keuangan (BPK) melalui audit eksternal yang menghasilkan temuan-temuan terkait Sistem Pengendalian Intern (SPI) dan kepatuhan terhadap peraturan perundang-undangan. Temuan BPK dapat dipandang sebagai indikator tata kelola yang bersifat ex-post, melengkapi MCP yang lebih bersifat ex-ante. Interaksi antara dua instrumen pengawasan ini—internal (MCP) dan eksternal (BPK)—membentuk apa yang dapat disebut sebagai dual governance framework, namun belum banyak diteliti secara empiris."
)

add_para(doc,
    "Selain itu, KPK melakukan revisi kerangka indikator MCP dari label \"Perizinan\" pada 2023 menjadi \"Pelayanan Publik\" pada 2024-2025, dengan penambahan sub-indikator Standar Pelayanan, Survei Kepuasan Masyarakat, dan Layanan Publik Berintegritas. Perubahan ini memberi peluang untuk menguji apakah penyempurnaan kerangka indikator tata kelola memberi dampak yang berbeda terhadap respon investor."
)

add_para(doc,
    "Berdasarkan latar belakang tersebut, penelitian ini bertujuan menganalisis pengaruh komponen MCP terhadap investasi PMDN di Provinsi Jawa Tengah dengan pendekatan dekomposisi per sub-area, mempertimbangkan moderasi temuan BPK, dan menguji konsistensi temuan terhadap perubahan kerangka indikator. Penelitian ini diharapkan memberikan kontribusi baik secara metodologis (dekomposisi MCP, dual governance framework) maupun praktis (rekomendasi prioritas reformasi tata kelola untuk pemerintah daerah)."
)

# ============ 1.2 Rumusan Masalah ============
add_heading_custom(doc, "1.2 Rumusan Masalah", level=1)
add_para(doc,
    "Berdasarkan latar belakang yang telah diuraikan, rumusan masalah penelitian ini adalah:",
    indent_first=False
)

rms = [
    "Komponen Monitoring Center for Prevention (MCP) manakah yang berpengaruh signifikan terhadap nilai investasi Penanaman Modal Dalam Negeri (PMDN) di kabupaten/kota Provinsi Jawa Tengah periode 2023-2025?",
    "Apakah jumlah temuan Badan Pemeriksa Keuangan (BPK) memoderasi hubungan antara komponen MCP dan investasi PMDN di kabupaten/kota Provinsi Jawa Tengah?",
    "Apakah perubahan kerangka indikator MCP dari \"Perizinan\" (2023) menjadi \"Pelayanan Publik\" (2024-2025) menghasilkan pengaruh yang berbeda terhadap investasi PMDN?",
]
for i, q in enumerate(rms, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i}.   {q}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# ============ 1.3 Tujuan Penelitian ============
add_heading_custom(doc, "1.3 Tujuan Penelitian", level=1)
add_para(doc,
    "Sesuai dengan rumusan masalah, tujuan penelitian ini adalah:",
    indent_first=False
)

tujuans = [
    "Mengidentifikasi komponen MCP yang berpengaruh signifikan terhadap nilai investasi PMDN di kabupaten/kota Provinsi Jawa Tengah periode 2023-2025.",
    "Menganalisis peran jumlah temuan BPK sebagai variabel moderasi dalam hubungan antara komponen MCP dan investasi PMDN.",
    "Menguji konsistensi pengaruh komponen MCP terhadap investasi PMDN sebelum dan sesudah perubahan kerangka indikator dari \"Perizinan\" menjadi \"Pelayanan Publik\".",
]
for i, q in enumerate(tujuans, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i}.   {q}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# ============ 1.4 Manfaat Penelitian ============
add_heading_custom(doc, "1.4 Manfaat Penelitian", level=1)

add_heading_custom(doc, "1.4.1 Manfaat Teoretis", level=2)
manfaat_teori = [
    "Memberikan kontribusi metodologis berupa pendekatan dekomposisi MCP per sub-area, sebagai alternatif penggunaan skor agregat yang dominan di literatur sebelumnya.",
    "Memperkenalkan kerangka dual governance MCP-BPK sebagai mekanisme pengawasan internal-eksternal yang saling melengkapi dalam mempengaruhi investasi.",
    "Memperkaya literatur empiris tentang pengaruh tata kelola pemerintahan terhadap investasi di tingkat sub-nasional, khususnya pada level kabupaten/kota.",
]
for i, q in enumerate(manfaat_teori, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i}.   {q}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

add_heading_custom(doc, "1.4.2 Manfaat Praktis", level=2)
manfaat_praktis = [
    "Memberikan rekomendasi kebijakan kepada pemerintah daerah Jawa Tengah mengenai prioritas reformasi tata kelola yang paling efektif dalam mendorong investasi PMDN.",
    "Membantu KPK dalam mengevaluasi efektivitas kerangka indikator MCP, khususnya dampak revisi indikator pada periode 2024-2025.",
    "Memberikan informasi kepada calon investor mengenai kondisi tata kelola pelayanan publik di Jawa Tengah sebagai bahan pertimbangan keputusan investasi.",
]
for i, q in enumerate(manfaat_praktis, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i}.   {q}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# ============ 1.5 Sistematika Penulisan ============
add_heading_custom(doc, "1.5 Sistematika Penulisan", level=1)
add_para(doc,
    "Penelitian ini disusun dalam lima bab dengan rincian sebagai berikut:",
    indent_first=False
)

sis = [
    ("BAB I PENDAHULUAN", "menjelaskan latar belakang penelitian, rumusan masalah, tujuan, manfaat, dan sistematika penulisan."),
    ("BAB II TINJAUAN PUSTAKA", "menguraikan landasan teori (teori biaya transaksi, teori kelembagaan, teori agen-prinsipal), telaah penelitian terdahulu, kerangka konseptual, dan hipotesis penelitian."),
    ("BAB III METODE PENELITIAN", "memaparkan jenis dan pendekatan penelitian, populasi dan sampel, jenis dan sumber data, definisi operasional variabel, model regresi, serta teknik analisis data."),
    ("BAB IV HASIL DAN PEMBAHASAN", "menyajikan statistik deskriptif, hasil pemilihan model panel, hasil estimasi regresi, uji robustness, serta pembahasan temuan dikaitkan dengan teori dan literatur."),
    ("BAB V PENUTUP", "berisi kesimpulan, implikasi kebijakan, keterbatasan penelitian, dan saran untuk penelitian selanjutnya."),
]
for h, d in sis:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(h)
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(f", {d}")
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

doc.add_page_break()

# ============ BAB II TINJAUAN PUSTAKA ============
add_heading_custom(doc, "BAB II", level=0, center=True)
add_heading_custom(doc, "TINJAUAN PUSTAKA", level=0, center=True)

# 2.1 Landasan Teori
add_heading_custom(doc, "2.1 Landasan Teori", level=1)

add_heading_custom(doc, "2.1.1 Teori Biaya Transaksi", level=2)
add_para(doc,
    "Teori biaya transaksi (Coase, 1937; Williamson, 2002) menyatakan bahwa kegiatan ekonomi tidak hanya melibatkan biaya produksi (production cost) tetapi juga biaya transaksi (transaction cost) yang muncul dari proses pencarian informasi, negosiasi, kontrak, monitoring, dan enforcement. Investor swasta memilih lokasi investasi yang dapat meminimalkan biaya transaksi, salah satunya melalui kemudahan birokrasi dan transparansi pelayanan."
)
add_para(doc,
    "Korupsi dan birokrasi yang tidak efisien meningkatkan biaya transaksi karena investor harus mengeluarkan biaya tidak resmi (informal cost), menghadapi ketidakpastian, dan mengalami delay perizinan. Mauro (1995) menemukan bahwa korupsi yang tinggi menurunkan pertumbuhan investasi, sementara Mengistu dan Adhikary (2011) menyatakan bahwa investasi sangat elastis terhadap biaya transaksi terkait tata kelola. Implikasi bagi penelitian ini adalah bahwa upaya pencegahan korupsi yang langsung berdampak pada pengurangan biaya transaksi—khususnya pada sub-area perizinan dan pelayanan publik—diharapkan paling responsif terhadap keputusan investasi."
)

add_heading_custom(doc, "2.1.2 Teori Kelembagaan", level=2)
add_para(doc,
    "North (1991) mendefinisikan institusi sebagai aturan main yang dirancang untuk mengurangi ketidakpastian dalam interaksi ekonomi. Kualitas institusi suatu negara/daerah—tercermin dari supremasi hukum, transparansi birokrasi, dan akuntabilitas pemerintah—merupakan determinan penting investasi. Aysan dkk. (2007) dan Globerman dkk. (2006) membuktikan bahwa kualitas tata kelola berpengaruh positif signifikan terhadap arus investasi, terutama di negara berkembang."
)
add_para(doc,
    "Dalam konteks Indonesia, Monitoring Center for Prevention (MCP) yang dikembangkan KPK merupakan instrumen pengukuran kualitas institusi tata kelola di tingkat pemerintah daerah. Skor MCP yang tinggi mencerminkan kelembagaan yang lebih kredibel, sehingga secara teoretis akan menarik lebih banyak investasi."
)

add_heading_custom(doc, "2.1.3 Teori Agen-Prinsipal", level=2)
add_para(doc,
    "Teori agen-prinsipal (Klitgaard, 1988; Camargo, 2011) memandang hubungan antara masyarakat (prinsipal) dan pemerintah (agen) sebagai relasi yang rawan asimetri informasi. Agen memiliki akses informasi yang lebih banyak dibandingkan prinsipal, sehingga muncul potensi praktik korupsi. Untuk mengatasi asimetri informasi tersebut diperlukan mekanisme akuntabilitas yang melibatkan: delegasi kewenangan, pembiayaan transparan, enforcement, informasi kinerja, dan monitoring."
)
add_para(doc,
    "MCP berperan sebagai mekanisme monitoring ex-ante yang menilai upaya pencegahan korupsi pemerintah daerah, sementara temuan BPK berfungsi sebagai monitoring ex-post yang memeriksa kepatuhan dan integritas pengelolaan keuangan. Keduanya bersifat saling melengkapi (komplementer) dalam membentuk kerangka tata kelola yang akuntabel, dan investor swasta dapat memandang kombinasi keduanya sebagai sinyal kualitas pemerintahan daerah."
)

# 2.2 Penelitian Terdahulu
add_heading_custom(doc, "2.2 Penelitian Terdahulu", level=1)

add_para(doc,
    "Beberapa penelitian terdahulu yang relevan dengan topik ini disajikan dalam Tabel 2.1.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Peneliti", "Fokus Penelitian", "Temuan Utama"],
    [
        ["Tua & Mahi (2022)", "Pengaruh MCP terhadap investasi swasta di 507 kab/kota Indonesia, 2018-2020", "MCP berpengaruh positif signifikan terhadap investasi swasta (β=0,040; p=0,015)"],
        ["Almakkiyah (2026)", "Pengaruh MCP terhadap investasi di 10 kab/kota NTB, 2019-2024 (panel REM)", "Hubungan nonlinier (threshold consistency); AKOR² positif signifikan; AGLO, HCAP, PAK signifikan"],
        ["Mauro (1995)", "Pengaruh korupsi terhadap pertumbuhan ekonomi lintas negara", "Korupsi tinggi menurunkan investasi melalui peningkatan biaya transaksi"],
        ["Aysan dkk. (2007)", "Tata kelola dan FDI di MENA region", "Kualitas tata kelola berpengaruh positif terhadap arus FDI"],
        ["Mengistu & Adhikary (2011)", "Tata kelola dan FDI di Asia, 1996-2007", "Investasi elastis terhadap biaya transaksi institusional"],
        ["Peres dkk. (2018)", "Pencegahan korupsi dan FDI di developing countries", "Hasil tidak signifikan karena lemahnya rule of law"],
    ],
    caption="Tabel 2.1 Ringkasan Penelitian Terdahulu"
)

doc.add_paragraph()
add_para(doc,
    "Berdasarkan telaah penelitian terdahulu, terdapat beberapa research gap yang menjadi celah penelitian ini:",
    indent_first=False
)
gaps = [
    "Penelitian terdahulu menggunakan skor MCP secara agregat sehingga belum dapat mengidentifikasi sub-area MCP yang menjadi penggerak utama hubungan dengan investasi.",
    "Belum ada penelitian yang menempatkan temuan BPK sebagai variabel moderasi dalam hubungan MCP-investasi (kerangka dual governance).",
    "Studi pada level provinsi spesifik dengan karakter heterogen seperti Jawa Tengah masih terbatas; mayoritas penelitian beroperasi di level nasional atau provinsi dengan jumlah kab/kota kecil.",
    "Belum ada penelitian yang mengakomodasi perubahan kerangka indikator MCP secara empiris (uji structural break).",
]
for i, q in enumerate(gaps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i}.   {q}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# 2.3 Kerangka Konseptual
add_heading_custom(doc, "2.3 Kerangka Konseptual", level=1)

add_para(doc,
    "Berdasarkan landasan teori dan penelitian terdahulu, kerangka konseptual penelitian ini dirumuskan sebagai berikut. Komponen MCP yang merepresentasikan kualitas tata kelola pencegahan korupsi internal pemerintah daerah diharapkan berpengaruh positif terhadap investasi PMDN. Dari tujuh sub-area MCP yang diuji, sub-area Perizinan/Pelayanan Publik diprediksi paling berpengaruh karena langsung berinteraksi dengan investor (sesuai teori biaya transaksi). Jumlah temuan BPK berperan sebagai variabel moderasi yang dapat memperkuat atau melemahkan pengaruh MCP terhadap investasi (kerangka dual governance). Variabel jumlah penduduk digunakan sebagai variabel kontrol untuk merepresentasikan ukuran pasar potensial daerah."
)

doc.add_paragraph()
# Visual kerangka konseptual sederhana
kerangka_text = """
  ┌──────────────────────┐                ┌──────────────────────┐
  │  KOMPONEN MCP (X)    │                │   TEMUAN BPK (M)     │
  │  ─────────────────   │                │   ─────────────────  │
  │  • MCP_PERANG        │   moderasi     │   • TOT_TEMUAN       │
  │  • MCP_IZIN          │ ◄─────────────►│   • TEM_SPI          │
  │  • MCP_MASN          │                │   • TEM_KEP          │
  │  • MCP_APIP          │                └──────────┬───────────┘
  │  • MCP_MASET         │                           │
  │  • MCP_OPD           │                           │
  │  • MCP_PBJ           │                           │
  └──────────┬───────────┘                           │
             │                                       │
             │            pengaruh utama             │
             └─────────────────┬─────────────────────┘
                               │
                               ▼
                  ┌────────────────────────┐
                  │  INVESTASI PMDN (Y)    │
                  │  log(INVEST + 1)       │
                  └────────────────────────┘
                               ▲
                               │
                  ┌────────────┴───────────┐
                  │   VARIABEL KONTROL     │
                  │  ─────────────────     │
                  │  log(PENDUDUK)         │
                  └────────────────────────┘
"""
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(kerangka_text)
run.font.name = "Courier New"
run.font.size = Pt(9)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gambar 2.1 Kerangka Konseptual Penelitian")
r.bold = True
r.italic = True
r.font.size = Pt(11)
r.font.name = "Times New Roman"

# 2.4 Hipotesis
add_heading_custom(doc, "2.4 Hipotesis Penelitian", level=1)

add_para(doc,
    "Berdasarkan kerangka konseptual, hipotesis penelitian dirumuskan sebagai berikut:",
    indent_first=False
)

hipos = [
    ("H1", "Komponen MCP Perizinan/Pelayanan Publik berpengaruh positif signifikan terhadap investasi PMDN di kabupaten/kota Provinsi Jawa Tengah."),
    ("H2", "Komponen MCP Perencanaan & Penganggaran, Pengadaan Barang & Jasa, Pengawasan APIP, Manajemen ASN, Optimalisasi Pajak Daerah, dan Manajemen Aset berpengaruh positif terhadap investasi PMDN dengan magnitudo yang berbeda-beda."),
    ("H3", "Jumlah temuan BPK memoderasi hubungan antara komponen MCP dan investasi PMDN."),
    ("H4", "Pengaruh komponen MCP terhadap investasi PMDN berbeda secara signifikan antara periode dengan kerangka indikator \"Perizinan\" (2023) dan \"Pelayanan Publik\" (2024-2025)."),
]

for code, h in hipos:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{code}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(h)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

doc.add_page_break()

# ============ BAB III METODE PENELITIAN ============
add_heading_custom(doc, "BAB III", level=0, center=True)
add_heading_custom(doc, "METODE PENELITIAN", level=0, center=True)

# 3.1 Jenis Penelitian
add_heading_custom(doc, "3.1 Jenis dan Pendekatan Penelitian", level=1)
add_para(doc,
    "Penelitian ini menggunakan pendekatan kuantitatif dengan metode regresi data panel. Pendekatan kuantitatif dipilih karena penelitian ini berfokus pada pengujian hubungan antar variabel secara empiris menggunakan data numerik dan analisis statistik. Metode regresi data panel digunakan karena data penelitian mencakup gabungan dimensi cross-section (35 kabupaten/kota) dan time series (3 tahun, 2023-2025), sehingga memungkinkan analisis variasi antar daerah dan antar waktu secara simultan."
)

# 3.2 Lokasi & Waktu
add_heading_custom(doc, "3.2 Lokasi dan Waktu Penelitian", level=1)
add_para(doc,
    "Penelitian dilaksanakan pada seluruh kabupaten/kota di Provinsi Jawa Tengah yang berjumlah 35 unit, terdiri atas 29 kabupaten dan 6 kota (Semarang, Surakarta, Salatiga, Pekalongan, Tegal, Magelang). Periode pengamatan adalah tahun 2023-2025 (tiga tahun) dengan pertimbangan bahwa pada periode tersebut KPK menggunakan kerangka indikator MCP berbasis risiko yang relatif konsisten secara konseptual, meskipun terdapat revisi label sub-area pada 2024."
)

# 3.3 Populasi & Sampel
add_heading_custom(doc, "3.3 Populasi dan Sampel", level=1)
add_para(doc,
    "Populasi penelitian adalah seluruh kabupaten/kota di Provinsi Jawa Tengah. Penelitian ini menggunakan teknik sensus (sampel jenuh) dengan mengambil seluruh anggota populasi sebagai sampel, yaitu 35 kabupaten/kota. Dengan periode pengamatan 3 tahun, total observasi yang digunakan adalah 105 (35 × 3) dengan struktur balanced panel."
)

# 3.4 Jenis & Sumber Data
add_heading_custom(doc, "3.4 Jenis dan Sumber Data", level=1)
add_para(doc,
    "Jenis data yang digunakan adalah data sekunder yang diperoleh dari publikasi resmi lembaga-lembaga terkait. Pengumpulan data dilakukan melalui studi dokumentasi dengan menghimpun, mencatat, dan mengkaji dokumen serta publikasi resmi yang relevan. Sumber data dirinci pada Tabel 3.1."
)

doc.add_paragraph()
add_table(doc,
    ["No", "Variabel", "Sumber Data"],
    [
        ["1", "Nilai Realisasi Investasi PMDN", "BPS Jawa Tengah / DPMPTSP Provinsi Jawa Tengah"],
        ["2", "Skor MCP per sub-area", "Komisi Pemberantasan Korupsi (KPK) - Dashboard MCP"],
        ["3", "Jumlah Temuan BPK (Total, SPI, Kepatuhan)", "Badan Pemeriksa Keuangan (BPK) - IHPS"],
        ["4", "Jumlah Penduduk", "Badan Pusat Statistik (BPS) Jawa Tengah"],
        ["5", "PDRB ADHK", "Badan Pusat Statistik (BPS) Jawa Tengah"],
    ],
    caption="Tabel 3.1 Sumber Data Penelitian"
)

# 3.5 Definisi Operasional Variabel
add_heading_custom(doc, "3.5 Definisi Operasional Variabel", level=1)
add_para(doc,
    "Variabel-variabel yang digunakan dalam penelitian ini didefinisikan secara operasional sebagaimana disajikan dalam Tabel 3.2.",
    indent_first=False
)

doc.add_paragraph()
add_table(doc,
    ["Variabel", "Simbol", "Definisi Operasional", "Indikator/Rumus", "Satuan"],
    [
        ["Investasi PMDN", "INVEST", "Nilai realisasi Penanaman Modal Dalam Negeri kab/kota per tahun", "log(INVEST + 1)", "log Rupiah"],
        ["MCP Perencanaan-Penganggaran", "MCP_PERANG", "Skor pencegahan korupsi sub-area Perencanaan & Penganggaran APBD", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP Perizinan/Pelayanan Publik", "MCP_IZIN", "Skor pencegahan korupsi sub-area pelayanan perizinan/publik", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP Pengadaan", "MCP_PBJ", "Skor pencegahan korupsi sub-area Pengadaan Barang & Jasa", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP APIP", "MCP_APIP", "Skor pencegahan korupsi sub-area Pengawasan APIP", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP Manajemen ASN", "MCP_MASN", "Skor pencegahan korupsi sub-area Manajemen ASN", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP Optimalisasi Pajak", "MCP_OPD", "Skor pencegahan korupsi sub-area Optimalisasi Pajak Daerah", "Skor MCP KPK", "Indeks 0-100"],
        ["MCP Manajemen Aset", "MCP_MASET", "Skor pencegahan korupsi sub-area Manajemen Aset", "Skor MCP KPK", "Indeks 0-100"],
        ["Total Temuan BPK", "TOT_TEMUAN", "Total jumlah temuan BPK pada LHP kab/kota per tahun", "Total temuan", "Unit"],
        ["Temuan SPI", "TEM_SPI", "Jumlah temuan Sistem Pengendalian Intern", "Total temuan SPI", "Unit"],
        ["Temuan Kepatuhan", "TEM_KEP", "Jumlah temuan Kepatuhan terhadap peraturan", "Total temuan kepatuhan", "Unit"],
        ["Jumlah Penduduk", "PENDUDUK", "Jumlah penduduk kab/kota per tahun", "log(PENDUDUK)", "log Jiwa"],
        ["Dummy Periode", "D_POST2024", "Dummy untuk periode setelah perubahan label MCP", "1 jika tahun ≥ 2024; 0 jika tahun = 2023", "0/1"],
    ],
    caption="Tabel 3.2 Definisi Operasional Variabel"
)
doc.add_paragraph()
add_para(doc,
    "Catatan: Variabel INVEST dan PENDUDUK ditransformasi logaritma natural untuk mengatasi distribusi yang skewed dan menstandardisasi skala koefisien. Penambahan +1 pada log(INVEST+1) mengantisipasi nilai investasi nol.",
    italic=True, size=11
)

# 3.6 Model Penelitian
add_heading_custom(doc, "3.6 Model Penelitian", level=1)

add_para(doc,
    "Penelitian ini menggunakan lima spesifikasi model regresi data panel yang dirumuskan sebagai berikut:",
    indent_first=False
)

# Model 1
add_para(doc, "Model 1: Model Utama (Baseline)", bold=True)
add_para(doc,
    "Model utama menguji pengaruh komponen MCP Perizinan/Pelayanan Publik—yang diasumsikan paling berpengaruh berdasarkan teori biaya transaksi—terhadap investasi PMDN dengan kontrol jumlah penduduk.",
    indent_first=False
)
add_equation(doc,
    "log(INVESTᵢₜ + 1) = α + β₁·MCP_IZINᵢₜ + γ·log(PENDUDUKᵢₜ) + εᵢₜ",
    "...(3.1)"
)

# Model 2
doc.add_paragraph()
add_para(doc, "Model 2: Uji Structural Break (Perubahan Label MCP)", bold=True)
add_para(doc,
    "Model 2 menguji apakah perubahan kerangka indikator MCP dari \"Perizinan\" (2023) menjadi \"Pelayanan Publik\" (2024-2025) menghasilkan pengaruh yang berbeda terhadap investasi.",
    indent_first=False
)
add_equation(doc,
    "log(INVESTᵢₜ + 1) = α + β₁·MCP_IZINᵢₜ + β₂·D_POST2024ₜ "
    "+ β₃·(MCP_IZINᵢₜ × D_POST2024ₜ) + γ·log(PENDUDUKᵢₜ) + εᵢₜ",
    "...(3.2)"
)

# Model 3
doc.add_paragraph()
add_para(doc, "Model 3: Moderasi Temuan BPK", bold=True)
add_para(doc,
    "Model 3 menguji peran moderasi temuan BPK dalam hubungan MCP Perizinan dan investasi PMDN, dalam kerangka dual governance.",
    indent_first=False
)
add_equation(doc,
    "log(INVESTᵢₜ + 1) = α + β₁·MCP_IZINᵢₜ + β₂·TOT_TEMUANᵢₜ "
    "+ β₃·(MCP_IZINᵢₜ × TOT_TEMUANᵢₜ) + γ·log(PENDUDUKᵢₜ) + εᵢₜ",
    "...(3.3)"
)

# Model 4
doc.add_paragraph()
add_para(doc, "Model 4: Komponen MCP Lengkap (Identifikasi Driver)", bold=True)
add_para(doc,
    "Model 4 menguji seluruh komponen MCP secara bersamaan untuk mengidentifikasi sub-area yang paling dominan setelah dikontrol komponen lain.",
    indent_first=False
)
add_equation(doc,
    "log(INVESTᵢₜ + 1) = α + β₁·MCP_APIPᵢₜ + β₂·MCP_IZINᵢₜ + β₃·MCP_MASETᵢₜ + β₄·MCP_MASNᵢₜ "
    "+ β₅·MCP_OPDᵢₜ + β₆·MCP_PBJᵢₜ + β₇·MCP_PERANGᵢₜ + γ·log(PENDUDUKᵢₜ) + εᵢₜ",
    "...(3.4)"
)

# Model 5
doc.add_paragraph()
add_para(doc, "Model 5: Pemisahan Jenis Temuan BPK", bold=True)
add_para(doc,
    "Model 5 sebagai robustness check, memisahkan temuan BPK menjadi temuan SPI dan Kepatuhan untuk melihat apakah jenis temuan yang berbeda memberi pengaruh berbeda.",
    indent_first=False
)
add_equation(doc,
    "log(INVESTᵢₜ + 1) = α + β₁·MCP_IZINᵢₜ + β₂·TEM_SPIᵢₜ + β₃·TEM_KEPᵢₜ "
    "+ γ·log(PENDUDUKᵢₜ) + εᵢₜ",
    "...(3.5)"
)

doc.add_paragraph()
add_para(doc,
    "Keterangan: i = kabupaten/kota (1, 2, ..., 35); t = tahun (2023, 2024, 2025); α = konstanta; βⱼ = koefisien variabel utama; γ = koefisien variabel kontrol; εᵢₜ = error term.",
    italic=True, size=11
)

# 3.7 Teknik Analisis Data
add_heading_custom(doc, "3.7 Teknik Analisis Data", level=1)

add_heading_custom(doc, "3.7.1 Analisis Statistik Deskriptif", level=2)
add_para(doc,
    "Statistik deskriptif digunakan untuk menggambarkan karakteristik data masing-masing variabel, meliputi nilai rata-rata, median, nilai maksimum, nilai minimum, dan standar deviasi. Analisis ini bertujuan memberikan gambaran awal mengenai pola dan sebaran data sebelum dilakukan analisis regresi."
)

add_heading_custom(doc, "3.7.2 Pemilihan Model Regresi Panel", level=2)
add_para(doc,
    "Pemilihan antara Common Effect Model (CEM), Fixed Effect Model (FEM), dan Random Effect Model (REM) dilakukan melalui tiga pengujian formal:",
    indent_first=False
)

uji_panel = [
    ("Uji Chow", "Membandingkan CEM vs FEM. Hipotesis null: CEM lebih tepat. Jika probabilitas < 0,05, FEM dipilih."),
    ("Uji Hausman", "Membandingkan FEM vs REM. Hipotesis null: REM konsisten. Jika probabilitas < 0,05, FEM dipilih; jika ≥ 0,05, REM dipilih."),
    ("Uji Lagrange Multiplier (LM) Breusch-Pagan", "Membandingkan CEM vs REM. Hipotesis null: tidak ada random effect. Jika probabilitas ≥ 0,05, CEM dipilih; jika < 0,05, REM dipilih."),
]
for nm, desc in uji_panel:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"• {nm}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(desc)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

add_heading_custom(doc, "3.7.3 Uji Asumsi Klasik", level=2)
add_para(doc,
    "Uji asumsi klasik dilakukan untuk memastikan validitas estimasi, meliputi:",
    indent_first=False
)
asums = [
    ("Uji Normalitas", "menggunakan Jarque-Bera Test. Residual diharapkan terdistribusi normal."),
    ("Uji Multikolinearitas", "menggunakan Variance Inflation Factor (VIF). VIF < 10 menunjukkan tidak ada multikolinearitas serius."),
    ("Uji Heteroskedastisitas", "menggunakan Breusch-Pagan Test atau White Test."),
    ("Uji Autokorelasi", "menggunakan Durbin-Watson Statistic. Nilai DW mendekati 2 menunjukkan tidak ada autokorelasi."),
]
for nm, desc in asums:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"• {nm}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(desc)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

add_heading_custom(doc, "3.7.4 Pengujian Hipotesis", level=2)
add_para(doc,
    "Pengujian hipotesis dilakukan melalui tiga uji statistik:",
    indent_first=False
)
testing = [
    ("Uji t (Parsial)", "menguji signifikansi individual koefisien regresi pada taraf nyata α = 5% dan α = 10%."),
    ("Uji F (Simultan)", "menguji signifikansi seluruh variabel independen secara bersamaan terhadap variabel dependen."),
    ("Koefisien Determinasi (R² dan Adjusted R²)", "mengukur proporsi variasi variabel dependen yang dijelaskan oleh model."),
]
for nm, desc in testing:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"• {nm}: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(desc)
    r2.font.size = Pt(12)
    r2.font.name = "Times New Roman"

add_heading_custom(doc, "3.7.5 Robust Standard Error", level=2)
add_para(doc,
    "Estimasi dilakukan dengan White cross-section robust standard error untuk mengantisipasi heteroskedastisitas dan autokorelasi antar-unit. Untuk menguji ketahanan hasil, dilakukan estimasi pembanding dengan dan tanpa robust SE. Mengingat keterbatasan periode (T=3) yang dapat menyebabkan reduced rank pada robust SE, hasil OLS standar tetap dilaporkan sebagai model utama dengan justifikasi statistik Durbin-Watson yang menunjukkan tidak adanya autokorelasi residual."
)

add_heading_custom(doc, "3.7.6 Software Analisis", level=2)
add_para(doc,
    "Analisis data dilakukan menggunakan software EViews 12 untuk estimasi regresi data panel, uji pemilihan model, dan uji asumsi klasik."
)

doc.add_page_break()

# ============ DAFTAR PUSTAKA SEMENTARA ============
add_heading_custom(doc, "DAFTAR PUSTAKA (SEMENTARA)", level=0, center=True)

refs = [
    "Almakkiyah, M. (2026). Corruption Prevention Efforts Against Investment: Empirical Evidence in West Nusa Tenggara Province. Academia Open, 11(1).",
    "Aysan, A. F., Nabli, M. K., & Véganzonès-Varoudakis, M. A. (2007). Governance institutions and private investment: An application to the Middle East and North Africa. The Developing Economies, 45(3), 339-377.",
    "Camargo, C. (2011). Public Sector Anti-Corruption Frameworks: Components for Accountable Governance. U4 Anti-Corruption Resource Centre.",
    "Coase, R. H. (1937). The Nature of the Firm. Economica, 4(16), 386-405.",
    "Globerman, S., Shapiro, D., & Tang, Y. (2006). Foreign Direct Investment in Emerging and Transition European Countries. Journal of Multinational Financial Management, 16(2), 197-220.",
    "Klitgaard, R. (1988). Controlling Corruption. University of California Press.",
    "Mauro, P. (1995). Corruption and Growth. Quarterly Journal of Economics, 110(3), 681-712.",
    "Mengistu, A. A., & Adhikary, B. K. (2011). Does good governance matter for FDI inflows? Evidence from Asian economies. Asia Pacific Business Review, 17(3), 281-299.",
    "North, D. C. (1991). Institutions. Journal of Economic Perspectives, 5(1), 97-112.",
    "Peres, M., Ameer, W., & Xu, H. (2018). The impact of institutional quality on foreign direct investment inflows: evidence for developed and developing countries. Economic Research-Ekonomska Istraživanja, 31(1), 626-644.",
    "Tua, M., & Mahi, B. R. (2022). Analysis of the effect of corruption prevention on private investment at the district/city level in Indonesia. Integritas: Jurnal Antikorupsi, 8(2), 247-258.",
    "Williamson, O. E. (2002). The Theory of the Firm as Governance Structure: From Choice to Contract. Journal of Economic Perspectives, 16(3), 171-195.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(ref)
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"

# Save
output_path = "/projects/sandbox/investmcp/Draft_Bab_1-3.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")

import os
print(f"File size: {os.path.getsize(output_path):,} bytes")
