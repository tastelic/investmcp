"""
Build BAB V - PENUTUP
Skripsi: Dekomposisi Komponen MCP terhadap Investasi PMDN di Jawa Tengah 2023-2025
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT


def add_para(doc, text, bold=False, italic=False, size=12, justify=True, indent_first=True, space_after=6):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def add_heading_custom(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    if level == 0:
        run.font.size = Pt(14)
    elif level == 1:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    return p


def add_numbered(doc, text, number):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{number}.   {text}")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    return p


# ============ BUILD ============
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

for section in doc.sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ============ BAB V ============
add_heading_custom(doc, "BAB V", level=0, center=True)
add_heading_custom(doc, "PENUTUP", level=0, center=True)

# 5.1 Kesimpulan
add_heading_custom(doc, "5.1 Kesimpulan", level=1)
add_para(doc,
    "Berdasarkan hasil analisis dan pembahasan pada bab sebelumnya, dapat "
    "diambil kesimpulan sebagai berikut:",
    indent_first=False
)

kesimpulan = [
    "Dari tujuh komponen Monitoring Center for Prevention (MCP) yang diuji, hanya "
    "MCP Perizinan/Pelayanan Publik yang berpengaruh positif signifikan terhadap "
    "investasi Penanaman Modal Dalam Negeri (PMDN) di kabupaten/kota Provinsi "
    "Jawa Tengah periode 2023-2025 (β=0,0216; p=0,054). Sub-area MCP lainnya "
    "(Perencanaan-Penganggaran, Pengadaan, APIP, Manajemen ASN, Optimalisasi "
    "Pajak Daerah, dan Manajemen Aset) tidak menunjukkan pengaruh signifikan, "
    "demikian pula skor total MCP. Hal ini menunjukkan bahwa kontribusi MCP "
    "terhadap investasi terkonsentrasi pada satu sub-area kunci yaitu Perizinan "
    "yang langsung berinteraksi dengan investor swasta.",

    "Jumlah temuan Badan Pemeriksa Keuangan (BPK) tidak terbukti memoderasi "
    "hubungan antara MCP Perizinan dan investasi PMDN, baik dalam bentuk total "
    "temuan, temuan Sistem Pengendalian Intern (SPI), maupun temuan Kepatuhan. "
    "Hal ini mengindikasikan bahwa kerangka dual governance yang menempatkan MCP "
    "sebagai pengawasan internal dan BPK sebagai pengawasan eksternal belum "
    "berfungsi efektif sebagai mekanisme yang saling memperkuat dalam mempengaruhi "
    "keputusan investasi.",

    "Pengaruh MCP Perizinan terhadap investasi PMDN BERBEDA secara signifikan "
    "antara periode dengan kerangka indikator \"Perizinan\" (2023) dan \"Pelayanan "
    "Publik\" (2024-2025). Pengaruh signifikan baru muncul setelah revisi kerangka "
    "indikator KPK, dengan efek total β=0,054 pada periode 2024-2025 dibandingkan "
    "β=0,006 (tidak signifikan) pada 2023. Penambahan sub-indikator Standar "
    "Pelayanan, Survei Kepuasan Masyarakat, dan Layanan Publik Berintegritas "
    "membuat skor MCP lebih sensitif terhadap pengalaman aktual pengguna pelayanan, "
    "termasuk investor.",

    "Variabel kontrol jumlah penduduk konsisten signifikan di seluruh model "
    "(β≈0,46; p<0,01), mengindikasikan bahwa faktor pasar potensial (demografi) "
    "tetap menjadi determinan utama investasi PMDN di Jawa Tengah, melebihi "
    "pengaruh kualitas tata kelola.",
]
for i, k in enumerate(kesimpulan, 1):
    add_numbered(doc, k, i)

# 5.2 Implikasi Kebijakan
add_heading_custom(doc, "5.2 Implikasi Kebijakan", level=1)
add_para(doc,
    "Berdasarkan kesimpulan tersebut, beberapa implikasi kebijakan dapat "
    "dirumuskan sebagai berikut:",
    indent_first=False
)

implikasi = [
    "Bagi pemerintah daerah Provinsi Jawa Tengah, prioritas reformasi tata kelola "
    "untuk peningkatan investasi sebaiknya difokuskan pada Pelayanan Terpadu Satu "
    "Pintu (PTSP) dan kualitas pelayanan publik secara komprehensif. Hal ini "
    "lebih efektif dibandingkan kebijakan tata kelola yang merata pada seluruh "
    "delapan sub-area MCP. Investasi pada perbaikan sistem perizinan online, "
    "transparansi tata ruang, standar pelayanan, dan responsivitas terhadap "
    "kepuasan masyarakat akan memberikan dampak langsung terhadap iklim investasi.",

    "Bagi Komisi Pemberantasan Korupsi (KPK), penyempurnaan kerangka indikator MCP "
    "dari \"Perizinan\" menjadi \"Pelayanan Publik\" pada 2024 terbukti efektif dalam "
    "meningkatkan responsivitas investor. KPK perlu mempertahankan dan terus "
    "menyempurnakan substansi sub-indikator pelayanan publik dengan tetap menjaga "
    "konsistensi konseptual antar tahun untuk memungkinkan komparabilitas data.",

    "Bagi Badan Pemeriksa Keuangan (BPK), perlu dipertimbangkan strategi "
    "diseminasi hasil pemeriksaan yang lebih luas dan tepat sasaran kepada "
    "calon investor sebagai sinyal kualitas tata kelola daerah, sehingga "
    "temuan BPK dapat berfungsi efektif sebagai bagian dari kerangka dual "
    "governance yang berdampak pada keputusan investasi.",

    "Bagi calon investor swasta, penilaian iklim investasi suatu daerah "
    "sebaiknya tidak hanya berdasarkan skor agregat tata kelola, tetapi juga "
    "memperhatikan kualitas spesifik sub-area Perizinan/Pelayanan Publik yang "
    "menjadi indikator paling responsif terhadap kemudahan berusaha.",
]
for i, im in enumerate(implikasi, 1):
    add_numbered(doc, im, i)

# 5.3 Keterbatasan Penelitian
add_heading_custom(doc, "5.3 Keterbatasan Penelitian", level=1)
add_para(doc,
    "Penelitian ini memiliki beberapa keterbatasan yang perlu diakui:",
    indent_first=False
)

keterbatasan = [
    "Periode pengamatan relatif pendek (T=3 tahun) sehingga membatasi power "
    "statistik dan kemampuan menangkap dinamika jangka panjang. Khususnya, "
    "estimasi robust standard error dengan metode White cross-section "
    "menghasilkan reduced rank covariance matrix karena keterbatasan jumlah "
    "periode. Periode dipilih untuk menjaga konsistensi kerangka indikator MCP, "
    "namun perlu pengujian lanjutan dengan periode lebih panjang ketika data "
    "tersedia.",

    "Variansi skor MCP relatif rendah karena banyak kabupaten/kota di Jawa Tengah "
    "telah mencapai skor mendekati 100, sehingga sensitivitas estimasi pengaruh "
    "MCP terhadap investasi terbatas. Hal ini menjelaskan mengapa hanya satu "
    "komponen (Perizinan) yang terdeteksi signifikan.",

    "Perubahan label sub-area MCP dari \"Perizinan\" (2023) menjadi \"Pelayanan "
    "Publik\" (2024-2025) merupakan keterbatasan komparabilitas, meskipun sub-"
    "indikator yang diukur sebagian besar serupa (kompatibilitas sekitar 64%). "
    "Penelitian telah mengakomodasi hal ini melalui uji structural break, namun "
    "keterbatasan tetap perlu diakui.",

    "Penelitian ini fokus pada Provinsi Jawa Tengah dengan karakteristik unik "
    "(rata-rata MCP tinggi, ekonomi heterogen Pantura-pegunungan), sehingga "
    "generalisasi temuan ke provinsi lain perlu dilakukan dengan hati-hati.",

    "Variabel kontrol yang digunakan terbatas pada faktor demografi (jumlah "
    "penduduk). Faktor lain seperti infrastruktur, modal manusia, kepastian "
    "hukum, dan stabilitas politik daerah belum diakomodasi dalam model.",
]
for i, kt in enumerate(keterbatasan, 1):
    add_numbered(doc, kt, i)

# 5.4 Saran Penelitian Selanjutnya
add_heading_custom(doc, "5.4 Saran untuk Penelitian Selanjutnya", level=1)
add_para(doc,
    "Berdasarkan keterbatasan tersebut, beberapa saran untuk penelitian "
    "selanjutnya:",
    indent_first=False
)

saran = [
    "Memperluas periode pengamatan ke 2026 dan tahun-tahun berikutnya ketika "
    "data tersedia, dengan tetap menjaga konsistensi kerangka indikator MCP "
    "untuk memungkinkan komparasi yang valid antar periode.",

    "Memperluas wilayah penelitian ke provinsi lain di Pulau Jawa atau "
    "tingkat nasional untuk menguji generalisasi temuan dan membandingkan "
    "pengaruh komponen MCP antar konteks regional yang berbeda.",

    "Menambahkan variabel kontrol struktural lain seperti infrastruktur "
    "(panjang jalan, akses listrik), modal manusia (rata-rata lama sekolah), "
    "tingkat partisipasi angkatan kerja, dan PDRB per kapita untuk meningkatkan "
    "daya jelas model.",

    "Mengeksplorasi penggunaan data sub-indikator mentah dari KPK (bukan skor "
    "agregat 0-100) untuk meningkatkan variansi variabel MCP dan memperkaya "
    "informasi yang dianalisis.",

    "Mengembangkan model yang membedakan jenis investasi (manufaktur, jasa, "
    "perdagangan, properti) untuk menguji apakah sensitivitas terhadap MCP "
    "berbeda antar sektor ekonomi.",

    "Melengkapi kerangka dual governance dengan variabel pengawasan eksternal "
    "lain seperti opini BPK (WTP/WDP), indeks Sistem Akuntabilitas Kinerja "
    "Instansi Pemerintah (SAKIP), atau indeks tata kelola pemerintahan daerah "
    "yang lebih komprehensif.",
]
for i, s in enumerate(saran, 1):
    add_numbered(doc, s, i)

# Save
output_path = "/projects/sandbox/investmcp/Draft_Bab_5.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")

import os
print(f"File size: {os.path.getsize(output_path):,} bytes")
